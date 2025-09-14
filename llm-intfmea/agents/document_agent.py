"""
Combines data extraction pipeline with LlamaIndex RAG and advanced FMEA intelligence
"""

import io
import os
import re
import json
import time
import regex as re2
import requests
import pdfplumber
import docx
import tempfile
import logging
from dataclasses import dataclass
from typing import List, Dict, Tuple, Optional, Union, Any
from pathlib import Path
from datetime import datetime

import numpy as np
import pandas as pd
from rapidfuzz import fuzz
from sklearn.feature_extraction.text import TfidfVectorizer

# LlamaIndex imports
from llama_index.core import Settings
from llama_index.core import StorageContext
from llama_index.core import SimpleDirectoryReader
from llama_index.core.indices.vector_store import VectorStoreIndex
from llama_index.core.node_parser import SentenceSplitter
from llama_index.retrievers.bm25 import BM25Retriever
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.retrievers import QueryFusionRetriever
from llama_index.core.schema import Document
from llama_index.core.callbacks import CallbackManager, LlamaDebugHandler

# Vector store and embeddings
from llama_index.vector_stores.faiss import FaissVectorStore
from llama_index.embeddings.huggingface import HuggingFaceEmbedding

# LLM integrations
from llama_index.llms.ollama import Ollama
from llama_index.llms.openai import OpenAI

# Pydantic for structured output
from pydantic import BaseModel, Field
import faiss

try:
    import spacy
    from spacy.lang.en import English
    _SPACY_OK = True
except Exception:
    _SPACY_OK = False
    spacy = None

try:
    from langchain_community.embeddings import HuggingFaceEmbeddings as LC_HFEmbeddings
    from llama_index.embeddings.langchain import LangchainEmbedding
    _LANGCHAIN_OK = True
except Exception:
    _LANGCHAIN_OK = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- Global variables for spaCy and embeddings ---
_NLP = None
_SENT_EMB = None

def get_nlp():
    global _NLP
    if _NLP is not None:
        return _NLP
    if not _SPACY_OK:
        return None
    try:
        _NLP = spacy.load("en_core_web_sm")
    except Exception:
        try:
            from spacy.cli import download
            download("en_core_web_sm")
            _NLP = spacy.load("en_core_web_sm")
        except Exception:
            _NLP = English()
            _NLP.add_pipe("sentencizer")
    return _NLP

def get_sentence_embedder():
    global _SENT_EMB
    if _SENT_EMB is not None:
        return _SENT_EMB
    try:
        from sentence_transformers import SentenceTransformer
        _SENT_EMB = SentenceTransformer("all-MiniLM-L6-v2")
    except Exception:
        _SENT_EMB = None
    return _SENT_EMB

# --- Domain seed lexicons ---
DEFAULT_SUBSYSTEM_SEEDS = {
    "lubrication system", "hydraulic system", "cooling system", "coolant system",
    "pathpilot interface", "control system", "plc", "electrical cabinet",
    "spindle drive", "gearbox", "pneumatic system", "safety interlock",
}
DEFAULT_COMPONENT_SEEDS = {
    "spindle motor", "coolant pump", "tool holder", "bearing", "seal",
    "pressure sensor", "limit switch", "relay", "vfd", "belt", "filter",
    "nozzle", "fuse", "encoder", "actuator", "valve", "hose", "manifold",
}

# --- FMEA Intelligence from main2.py ---
class FMEAIntelligence:
    def __init__(self):
        # Severity indicators and their numerical mappings
        self.severity_keywords = {
            'high': {
                'keywords': ['fire', 'explosion', 'injury', 'death', 'fatality', 'catastrophic',
                             'major damage', 'complete failure', 'total shutdown', 'hazardous',
                             'toxic', 'environmental damage', 'regulatory violation'],
                'score': 8 - 10
            },
            'medium': {
                'keywords': ['downtime', 'production loss', 'quality issue', 'malfunction',
                             'partial failure', 'degraded performance', 'minor injury',
                             'rework required', 'customer complaint', 'delayed delivery'],
                'score': 4 - 7
            },
            'low': {
                'keywords': ['minor issue', 'cosmetic', 'slight reduction', 'negligible',
                             'barely noticeable', 'no safety impact', 'minor inconvenience'],
                'score': 1 - 3
            }
        }

        # Component-specific failure modes and their characteristics
        self.component_failure_db = {
            'motor': {
                'failure_modes': ['Overheating', 'Bearing failure', 'Electrical insulation breakdown',
                                  'Rotor imbalance'],
                'causes': ['Poor ventilation', 'Bearing wear', 'Moisture ingress', 'Manufacturing defects'],
                'effects': ['Complete motor failure', 'Reduced efficiency', 'Vibration and noise', 'System shutdown'],
                'high_risk_indicators': ['electrical', 'fire', 'explosion']
            },
            'pump': {
                'failure_modes': ['Cavitation', 'Seal leakage', 'Impeller wear', 'Clogging'],
                'causes': ['Insufficient NPSH', 'Seal deterioration', 'Abrasive particles', 'Foreign objects'],
                'effects': ['Flow reduction', 'Fluid contamination', 'Energy loss', 'System pressure drop'],
                'high_risk_indicators': ['leak', 'contamination', 'pressure']
            },
            'bearing': {
                'failure_modes': ['Fatigue spalling', 'Contamination', 'Misalignment', 'Lubrication failure'],
                'causes': ['Cyclic loading', 'Ingress of particles', 'Installation errors', 'Lubricant degradation'],
                'effects': ['Increased friction', 'Noise and vibration', 'Heat generation', 'Seizure'],
                'high_risk_indicators': ['seizure', 'fire', 'catastrophic']
            },
            'sensor': {
                'failure_modes': ['Drift', 'Signal loss', 'False readings', 'Physical damage'],
                'causes': ['Environmental conditions', 'Electrical interference', 'Calibration errors',
                           'Impact damage'],
                'effects': ['Incorrect measurements', 'Control system errors', 'Process instability',
                            'Safety system failure'],
                'high_risk_indicators': ['safety', 'control', 'critical']
            },
            'valve': {
                'failure_modes': ['Stuck closed', 'Stuck open', 'Internal leakage', 'Seat damage'],
                'causes': ['Corrosion', 'Debris accumulation', 'Actuator failure', 'Excessive pressure'],
                'effects': ['Flow blockage', 'Uncontrolled flow', 'Pressure loss', 'System imbalance'],
                'high_risk_indicators': ['pressure', 'uncontrolled', 'blockage']
            }
        }

        # Detailed recommendation templates
        self.recommendation_templates = {
            'high_severity': {
                'inspection': "Implement daily visual inspections and weekly detailed inspections using {inspection_method}. Document all findings in maintenance logs with photographic evidence.",
                'monitoring': "Install continuous monitoring systems with real-time alerts. Set up {monitoring_type} with alarm thresholds at 80% of safe operating limits.",
                'maintenance': "Establish preventive maintenance schedule every {frequency} with complete {component} replacement. Use OEM parts and certified technicians only.",
                'training': "Provide specialized safety training for all personnel handling this component. Include emergency response procedures and hazard recognition.",
                'backup': "Install redundant {component} systems with automatic failover capability. Ensure backup systems are regularly tested and maintained."
            },
            'medium_severity': {
                'inspection': "Conduct weekly visual inspections and monthly detailed inspections. Use {inspection_method} to assess component condition and performance.",
                'monitoring': "Implement condition-based monitoring using {monitoring_type}. Schedule monthly data analysis and trending.",
                'maintenance': "Perform preventive maintenance every {frequency}. Include lubrication, adjustment, and wear part replacement as needed.",
                'optimization': "Analyze operating conditions and optimize parameters to reduce stress on {component}. Consider upgrading to more robust alternatives."
            },
            'low_severity': {
                'inspection': "Include in routine monthly inspections. Visual assessment and basic functional testing sufficient.",
                'maintenance': "Follow manufacturer's recommended maintenance schedule. Document all activities for trend analysis.",
                'monitoring': "Include in quarterly condition assessments. No special monitoring equipment required."
            }
        }

    def extract_severity_from_text(self, text: str) -> Tuple[str, int, str]:
        """Extract severity level from historical data text"""
        text_lower = text.lower()

        # Initialize scores
        severity_scores = {'high': 0, 'medium': 0, 'low': 0}
        found_indicators = []

        # Score based on keyword presence
        for level, data in self.severity_keywords.items():
            for keyword in data['keywords']:
                if keyword in text_lower:
                    severity_scores[level] += 1
                    found_indicators.append(keyword)

        # Determine severity level
        max_score = max(severity_scores.values())
        if max_score == 0:
            return 'medium', 5, 'No specific indicators found'

        severity_level = max(severity_scores.keys(), key=lambda k: severity_scores[k])

        # Convert to numerical score
        if severity_level == 'high':
            numerical_score = min(10, 7 + severity_scores['high'])
        elif severity_level == 'medium':
            numerical_score = min(7, max(4, 4 + severity_scores['medium']))
        else:
            numerical_score = min(4, max(1, 1 + severity_scores['low']))

        explanation = f"Found indicators: {', '.join(found_indicators[:3])}"
        return severity_level, numerical_score, explanation

    def calculate_occurrence(self, component: str, failure_mode: str, historical_text: str = "") -> Tuple[int, str]:
        """Calculate occurrence rating based on component type and historical data"""
        base_occurrence = 5  # Default medium occurrence

        # Component-specific occurrence adjustments
        component_lower = component.lower()

        # High-wear components
        if any(term in component_lower for term in ['bearing', 'seal', 'belt', 'filter']):
            base_occurrence += 2

        # Critical electrical components
        elif any(term in component_lower for term in ['motor', 'sensor', 'relay', 'encoder']):
            base_occurrence += 1

        # Robust components
        elif any(term in component_lower for term in ['valve', 'manifold', 'housing']):
            base_occurrence -= 1

        # Analyze historical data if provided
        if historical_text:
            failure_indicators = ['failure', 'fault', 'breakdown', 'malfunction', 'issue', 'problem']
            frequency_indicators = ['frequent', 'recurring', 'repeated', 'multiple', 'often']

            historical_lower = historical_text.lower()
            failure_count = sum(1 for indicator in failure_indicators if indicator in historical_lower)
            frequency_count = sum(1 for indicator in frequency_indicators if indicator in historical_lower)

            if frequency_count > 0:
                base_occurrence += 2
            elif failure_count > 2:
                base_occurrence += 1

        # Ensure within valid range
        occurrence = max(1, min(10, base_occurrence))

        explanation = f"Based on component type ({component}) and failure pattern analysis"
        return occurrence, explanation

    def calculate_detection(self, component: str, failure_mode: str, severity: int) -> Tuple[int, str]:
        """Calculate detection rating based on component visibility and monitoring"""
        component_lower = component.lower()
        failure_lower = failure_mode.lower()

        # Default detection based on component accessibility
        if any(term in component_lower for term in ['sensor', 'display', 'indicator', 'gauge']):
            base_detection = 2  # Easy to detect
        elif any(term in component_lower for term in ['motor', 'pump', 'valve']):
            base_detection = 4  # Moderate detection
        elif any(term in component_lower for term in ['bearing', 'seal', 'internal']):
            base_detection = 7  # Difficult to detect
        else:
            base_detection = 5  # Average detection

        # Adjust based on failure mode
        if any(term in failure_lower for term in ['noise', 'vibration', 'leak', 'overheating']):
            base_detection -= 2  # Observable symptoms
        elif any(term in failure_lower for term in ['fatigue', 'wear', 'internal']):
            base_detection += 2  # Hidden failures

        # Critical failures need better detection
        if severity >= 8:
            base_detection = max(1, base_detection - 2)

        detection = max(1, min(10, base_detection))

        explanation = f"Based on component accessibility and failure mode observability"
        return detection, explanation

    def generate_detailed_recommendations(self, component: str, failure_mode: str,
                                          severity: int, occurrence: int, detection: int) -> str:
        """Generate detailed, specific recommendations"""
        component_lower = component.lower()
        component_type = None

        # Identify component type
        for comp_type in self.component_failure_db.keys():
            if comp_type in component_lower:
                component_type = comp_type
                break

        # Determine severity category
        if severity >= 8:
            severity_cat = 'high_severity'
        elif severity >= 4:
            severity_cat = 'medium_severity'
        else:
            severity_cat = 'low_severity'

        recommendations = []

        # Get appropriate templates
        templates = self.recommendation_templates[severity_cat]

        # Generate inspection recommendations
        if 'inspection' in templates:
            inspection_method = self._get_inspection_method(component_type)
            rec = templates['inspection'].format(
                inspection_method=inspection_method,
                component=component
            )
            recommendations.append(f"INSPECTION: {rec}")

        # Generate monitoring recommendations
        if 'monitoring' in templates:
            monitoring_type = self._get_monitoring_type(component_type)
            rec = templates['monitoring'].format(
                monitoring_type=monitoring_type,
                component=component
            )
            recommendations.append(f"MONITORING: {rec}")

        # Generate maintenance recommendations
        if 'maintenance' in templates:
            frequency = self._get_maintenance_frequency(severity, occurrence)
            rec = templates['maintenance'].format(
                frequency=frequency,
                component=component
            )
            recommendations.append(f"MAINTENANCE: {rec}")

        # Add severity-specific recommendations
        if severity_cat == 'high_severity':
            if 'training' in templates:
                rec = templates['training'].format(component=component)
                recommendations.append(f"TRAINING: {rec}")
            if 'backup' in templates:
                rec = templates['backup'].format(component=component)
                recommendations.append(f"REDUNDANCY: {rec}")
        elif severity_cat == 'medium_severity' and 'optimization' in templates:
            rec = templates['optimization'].format(component=component)
            recommendations.append(f"OPTIMIZATION: {rec}")

        return " | ".join(recommendations)

    def _get_inspection_method(self, component_type: str) -> str:
        methods = {
            'motor': 'thermal imaging and vibration analysis',
            'pump': 'flow measurement and visual leak inspection',
            'bearing': 'vibration analysis and oil analysis',
            'sensor': 'calibration verification and signal testing',
            'valve': 'stroke testing and seat leakage testing'
        }
        return methods.get(component_type, 'visual inspection and functional testing')

    def _get_monitoring_type(self, component_type: str) -> str:
        monitoring = {
            'motor': 'current signature analysis and temperature monitoring',
            'pump': 'flow rate and pressure monitoring',
            'bearing': 'vibration and acoustic monitoring',
            'sensor': 'signal drift monitoring',
            'valve': 'position feedback and pressure differential monitoring'
        }
        return monitoring.get(component_type, 'condition-based monitoring')

    def _get_maintenance_frequency(self, severity: int, occurrence: int) -> str:
        if severity >= 8 or occurrence >= 8:
            return "2 weeks"
        elif severity >= 6 or occurrence >= 6:
            return "1 month"
        elif severity >= 4 or occurrence >= 4:
            return "3 months"
        else:
            return "6 months"

# --- Historical Data Analysis from main2.py ---
def analyze_historical_data(historical_text: str) -> Dict:
    """Analyze historical maintenance data to extract patterns and common issues"""

    if not historical_text:
        return {}

    # Parse the historical data (assuming tab/comma separated format)
    lines = historical_text.strip().split('\n')
    historical_records = []

    for line in lines[1:]:  # Skip header
        if '\t' in line:
            parts = line.split('\t', 1)
        elif ',' in line:
            parts = line.split(',', 1)
        else:
            continue

        if len(parts) >= 2:
            description = parts[0].strip()
            action = parts[1].strip()
            historical_records.append({
                'description': description,
                'action': action
            })

    # Component mapping - map real component names to extracted components
    component_mappings = {
        'start btn': ['button', 'start button', 'control button'],
        'brake solenoid': ['solenoid', 'brake', 'brake solenoid'],
        'spindle': ['spindle', 'spindle motor'],
        'fan': ['fan', 'cooling fan', 'spindle fan'],
        'vfd': ['vfd', 'drive', 'variable frequency drive'],
        'drawbar': ['drawbar', 'tool holder'],
        'spring': ['spring', 'drawbar spring'],
        'coolant': ['coolant', 'mist coolant', 'coolant pump'],
        'nozzle': ['nozzle', 'spray nozzle', 'mist nozzle'],
        'pump': ['pump', 'coolant pump', 'hydraulic pump'],
        'oiler': ['oiler', 'auto oiler', 'lubrication system'],
        'pathpilot': ['pathpilot', 'controller', 'cnc controller'],
        'axis': ['axis', 'x axis', 'y axis', 'z axis'],
        'limit switch': ['switch', 'limit switch', 'limit sensor'],
        'usb': ['usb', 'cable', 'connector'],
        'bearing': ['bearing', 'spindle bearing'],
        'gib': ['gib', 'way', 'slide'],
        'rail': ['rail', 'guide rail', 'linear rail']
    }

    # Group records by component
    component_issues = {}

    for record in historical_records:
        desc = record['description'].lower()
        action = record['action'].lower()

        # Find which component this record refers to
        matched_component = None
        for main_component, variants in component_mappings.items():
            if any(variant in desc for variant in variants):
                matched_component = main_component
                break

        if matched_component:
            if matched_component not in component_issues:
                component_issues[matched_component] = {
                    'descriptions': [],
                    'actions': [],
                    'frequency': 0
                }

            component_issues[matched_component]['descriptions'].append(record['description'])
            component_issues[matched_component]['actions'].append(record['action'])
            component_issues[matched_component]['frequency'] += 1

    return component_issues

def find_best_match_component(component_name: str, historical_analysis: Dict) -> Tuple[str, List[str], List[str], int]:
    """Find the best matching component from historical data"""

    component_lower = component_name.lower()
    best_match = None
    best_score = 0

    # Direct keyword matching
    for hist_component, data in historical_analysis.items():
        # Check if historical component name is in extracted component or vice versa
        if hist_component in component_lower or any(word in hist_component for word in component_lower.split()):
            score = len(hist_component) + data['frequency']  # Prefer longer matches and frequent issues
            if score > best_score:
                best_score = score
                best_match = hist_component

    if best_match:
        data = historical_analysis[best_match]
        return best_match, data['descriptions'], data['actions'], data['frequency']

    return None, [], [], 0

def generate_synthetic_issues(component: str, component_type: str = None) -> Tuple[List[str], List[str]]:
    """Generate realistic synthetic issues when no historical data is available"""

    # Enhanced synthetic issue database
    synthetic_db = {
        'motor': {
            'descriptions': [
                f"{component} overheating during operation. Temperature reaching 85°C+",
                f"{component} vibration increased. Bearing wear suspected",
                f"{component} current draw higher than normal. Possible winding issue",
                f"{component} not starting. LED indicator blinking fault code",
                f"{component} speed fluctuation under load. Drive parameters may need adjustment"
            ],
            'actions': [
                f"Clean {component} cooling fins and check ventilation pathways",
                f"Replace {component} bearings and check alignment",
                f"Test {component} windings with megger. Check for moisture ingress",
                f"Reset {component} drive and check error codes. Update parameters if needed",
                f"Calibrate {component} speed control and load parameters"
            ]
        },
        'pump': {
            'descriptions': [
                f"{component} flow rate decreased by 20%. Possible cavitation or wear",
                f"{component} seal leaking. Fluid contamination risk",
                f"{component} making unusual noise during operation",
                f"{component} pressure fluctuation. Check for air entrainment",
                f"{component} not priming properly. Suction line issue suspected"
            ],
            'actions': [
                f"Check {component} impeller clearance and replace if worn",
                f"Replace {component} mechanical seals and O-rings",
                f"Inspect {component} bearings and lubrication system",
                f"Check {component} suction line for leaks and proper priming",
                f"Clean {component} strainer and check inlet conditions"
            ]
        },
        'bearing': {
            'descriptions': [
                f"{component} noise increased during rotation. Possible spalling",
                f"{component} temperature rise detected. Lubrication issue suspected",
                f"{component} play detected during inspection. Wear progression",
                f"{component} seized during operation. Emergency stop activated",
                f"{component} contamination visible. Seal failure indicated"
            ],
            'actions': [
                f"Replace {component} and check shaft condition for damage",
                f"Relubricate {component} with specified grade and quantity",
                f"Monitor {component} condition and plan replacement schedule",
                f"Emergency {component} replacement and shaft inspection required",
                f"Replace {component} seals and clean housing thoroughly"
            ]
        },
        'sensor': {
            'descriptions': [
                f"{component} reading drift detected. Calibration out of spec",
                f"{component} signal intermittent. Connection issue suspected",
                f"{component} false triggering during operation",
                f"{component} no response to stimulus. Power or wiring fault",
                f"{component} accuracy degraded beyond tolerance limits"
            ],
            'actions': [
                f"Recalibrate {component} using certified reference standards",
                f"Check {component} wiring connections and shield integrity",
                f"Adjust {component} mounting position and sensitivity settings",
                f"Test {component} power supply and replace if faulty",
                f"Replace {component} and update calibration records"
            ]
        },
        'valve': {
            'descriptions': [
                f"{component} sticking in closed position. Debris or corrosion",
                f"{component} internal leakage detected. Seat damage possible",
                f"{component} actuator response slow. Air supply or linkage issue",
                f"{component} cannot achieve full stroke. Obstruction suspected",
                f"{component} chattering during operation. Pressure instability"
            ],
            'actions': [
                f"Disassemble {component} and clean internal surfaces",
                f"Replace {component} seats and sealing surfaces",
                f"Check {component} actuator air supply and linkage alignment",
                f"Inspect {component} stem and guide for binding or damage",
                f"Adjust {component} control parameters and check supply pressure"
            ]
        },
        'switch': {
            'descriptions': [
                f"{component} intermittent operation. Contact wear suspected",
                f"{component} false triggering. Vibration or EMI interference",
                f"{component} physical damage to actuator lever",
                f"{component} no continuity when actuated. Internal failure",
                f"{component} position drift causing premature activation"
            ],
            'actions': [
                f"Replace {component} and clean mounting area",
                f"Relocate {component} away from vibration sources and check shielding",
                f"Replace {component} actuator lever and check alignment",
                f"Replace entire {component} assembly and test operation",
                f"Readjust {component} position and lock mounting hardware"
            ]
        },
        'generic': {
            'descriptions': [
                f"{component} performance degradation observed during routine inspection",
                f"{component} unusual wear pattern detected. Operating conditions review needed",
                f"{component} maintenance interval exceeded. Preventive service required",
                f"{component} minor defect noted. Monitoring for progression",
                f"{component} efficiency decreased. Optimization potential identified"
            ],
            'actions': [
                f"Schedule detailed inspection of {component} and related systems",
                f"Review {component} operating parameters and adjust if necessary",
                f"Perform scheduled maintenance on {component} per manufacturer specs",
                f"Continue monitoring {component} condition and document changes",
                f"Analyze {component} performance data and implement improvements"
            ]
        }
    }

    # Determine component type if not provided
    if not component_type:
        component_lower = component.lower()
        if any(term in component_lower for term in ['motor', 'spindle']):
            component_type = 'motor'
        elif any(term in component_lower for term in ['pump']):
            component_type = 'pump'
        elif any(term in component_lower for term in ['bearing']):
            component_type = 'bearing'
        elif any(term in component_lower for term in ['sensor', 'switch']):
            component_type = 'sensor'
        elif any(term in component_lower for term in ['valve']):
            component_type = 'valve'
        elif any(term in component_lower for term in ['switch', 'limit']):
            component_type = 'switch'
        else:
            component_type = 'generic'

    if component_type in synthetic_db:
        return synthetic_db[component_type]['descriptions'], synthetic_db[component_type]['actions']
    else:
        return synthetic_db['generic']['descriptions'], synthetic_db['generic']['actions']

# --- Enhanced FMEA Generation from main2.py ---
def generate_intelligent_fmea(components_df: pd.DataFrame, historical_text: str = "") -> pd.DataFrame:
    """Generate FMEA table with historical data integration and intelligent fallbacks"""

    # Analyze historical data first
    historical_analysis = analyze_historical_data(historical_text)

    fmea_intelligence = FMEAIntelligence()

    fmea_columns = [
        "Component", "Potential Failure Mode", "Potential Causes", "Potential Effects",
        "Severity Level", "Severity (S)", "Severity Reasoning",
        "Occurrence (O)", "Occurrence Reasoning",
        "Detection (D)", "Detection Reasoning",
        "RPN", "Detailed Recommended Actions", "Data Source"
    ]

    fmea_data = []

    for _, row in components_df.iterrows():
        component = row["canonical"]
        component_lower = component.lower()

        # Try to find matching historical data
        matched_component, historical_descriptions, historical_actions, frequency = \
            find_best_match_component(component, historical_analysis)

        if matched_component and historical_descriptions:
            # Use historical data
            data_source = f"Historical data ({frequency} occurrences)"

            # Use up to 3 most common issues from historical data
            for i, (desc, action) in enumerate(zip(historical_descriptions[:3], historical_actions[:3])):

                # Extract severity from the specific description
                severity_level, severity_score, severity_reasoning = \
                    fmea_intelligence.extract_severity_from_text(desc)

                # Calculate occurrence based on frequency in historical data
                base_occurrence = min(10, max(1, int(frequency / 2) + 3))  # More frequent = higher occurrence
                occurrence = base_occurrence
                occurrence_reasoning = f"Based on {frequency} historical occurrences of similar issues"

                # Calculate detection based on the description
                detection, detection_reasoning = fmea_intelligence.calculate_detection(
                    component, desc, severity_score
                )

                # Use actual historical description as failure mode
                failure_mode = desc

                # Derive causes and effects from description
                if "leak" in desc.lower():
                    cause = "Seal degradation, fitting looseness, or component wear"
                    effect = "Fluid loss, contamination risk, reduced system performance"
                elif "intermittent" in desc.lower() or "takes" in desc.lower():
                    cause = "Contact wear, debris accumulation, or electrical connection issues"
                    effect = "Operational delays, potential system lockup, user frustration"
                elif "not" in desc.lower() and ("start" in desc.lower() or "spray" in desc.lower()):
                    cause = "Power supply failure, blockage, or component malfunction"
                    effect = "System shutdown, process interruption, potential overheating"
                elif "noise" in desc.lower() or "clicking" in desc.lower():
                    cause = "Mechanical wear, misalignment, or debris contamination"
                    effect = "Accelerated wear, potential seizure, quality degradation"
                elif "temperature" in desc.lower() or "temp" in desc.lower():
                    cause = "Cooling system failure, excessive load, or ventilation blockage"
                    effect = "Component damage, thermal protection activation, fire risk"
                else:
                    cause = "Component wear, environmental factors, or operational stress"
                    effect = "Performance degradation, potential failure, maintenance required"

                rpn = severity_score * occurrence * detection

                # Use historical action as part of recommendation
                detailed_actions = f"HISTORICAL ACTION: {action} | " + \
                                   fmea_intelligence.generate_detailed_recommendations(
                                       component, failure_mode, severity_score, occurrence, detection
                                   )

                fmea_data.append({
                    "Component": component,
                    "Potential Failure Mode": failure_mode,
                    "Potential Causes": cause,
                    "Potential Effects": effect,
                    "Severity Level": severity_level.title(),
                    "Severity (S)": severity_score,
                    "Severity Reasoning": severity_reasoning,
                    "Occurrence (O)": occurrence,
                    "Occurrence Reasoning": occurrence_reasoning,
                    "Detection (D)": detection,
                    "Detection Reasoning": detection_reasoning,
                    "RPN": rpn,
                    "Detailed Recommended Actions": detailed_actions,
                    "Data Source": data_source
                })

        else:
            # No historical data - use synthetic realistic data
            data_source = "Synthetic engineering knowledge"

            # Determine component type for synthetic data
            component_type = None
            for comp_type in fmea_intelligence.component_failure_db.keys():
                if comp_type in component_lower:
                    component_type = comp_type
                    break

            # Generate synthetic issues
            synthetic_descriptions, synthetic_actions = generate_synthetic_issues(component, component_type)

            # Use up to 2 synthetic issues per component
            for i, (desc, action) in enumerate(zip(synthetic_descriptions[:2], synthetic_actions[:2])):

                # Extract severity
                severity_level, severity_score, severity_reasoning = \
                    fmea_intelligence.extract_severity_from_text(desc)

                # Calculate occurrence (lower than historical since it's not observed)
                occurrence, occurrence_reasoning = fmea_intelligence.calculate_occurrence(
                    component, desc, ""
                )
                occurrence = max(1, occurrence - 1)  # Reduce by 1 since not historically observed
                occurrence_reasoning = f"Estimated based on component type and industry standards (no historical data)"

                # Calculate detection
                detection, detection_reasoning = fmea_intelligence.calculate_detection(
                    component, desc, severity_score
                )

                # Parse synthetic description for causes and effects
                if "overheating" in desc.lower():
                    cause = "Poor ventilation, excessive load, cooling system failure"
                    effect = "Component damage, thermal shutdown, potential fire hazard"
                elif "vibration" in desc.lower():
                    cause = "Imbalance, bearing wear, mounting looseness"
                    effect = "Accelerated wear, noise, potential mechanical failure"
                elif "leak" in desc.lower():
                    cause = "Seal failure, pressure fluctuation, material degradation"
                    effect = "Fluid loss, contamination, environmental impact"
                elif "noise" in desc.lower():
                    cause = "Mechanical wear, lubrication failure, foreign objects"
                    effect = "Component degradation, operational disturbance"
                else:
                    cause = "Normal wear, environmental conditions, operational stress"
                    effect = "Performance reduction, increased maintenance, potential failure"

                rpn = severity_score * occurrence * detection

                # Combine synthetic action with intelligent recommendations
                detailed_actions = f"PREVENTIVE ACTION: {action} | " + \
                                   fmea_intelligence.generate_detailed_recommendations(
                                       component, desc, severity_score, occurrence, detection
                                   )

                fmea_data.append({
                    "Component": component,
                    "Potential Failure Mode": desc,
                    "Potential Causes": cause,
                    "Potential Effects": effect,
                    "Severity Level": severity_level.title(),
                    "Severity (S)": severity_score,
                    "Severity Reasoning": severity_reasoning,
                    "Occurrence (O)": occurrence,
                    "Occurrence Reasoning": occurrence_reasoning,
                    "Detection (D)": detection,
                    "Detection Reasoning": detection_reasoning,
                    "RPN": rpn,
                    "Detailed Recommended Actions": detailed_actions,
                    "Data Source": data_source
                })

    return pd.DataFrame(fmea_data)

# --- File loading utilities ---
def load_pdf(file: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file)) as pdf:
        for page in pdf.pages:
            try:
                text_parts.append(page.extract_text() or "")
            except Exception:
                continue
    return "\n".join(text_parts)

def load_docx(file: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file)
        tmp.flush()
        path = tmp.name
    try:
        d = docx.Document(path)
        text = []
        for p in d.paragraphs:
            text.append(p.text)
        for table in d.tables:
            for row in table.rows:
                text.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(text)
    finally:
        try:
            os.remove(path)
        except Exception:
            pass

def load_txt(file: bytes) -> str:
    try:
        return file.decode("utf-8", errors="ignore")
    except Exception:
        return file.decode(errors="ignore")

def load_excel_csv(file: bytes, fname: str) -> Dict[str, List[str]]:
    seeds = {"subsystems": [], "components": []}
    try:
        if fname.lower().endswith(".csv"):
            df = pd.read_csv(io.BytesIO(file))
        else:
            df = pd.read_excel(io.BytesIO(file))
        for col in df.columns:
            if "subsystem" in col.lower():
                seeds["subsystems"].extend(df[col].dropna().astype(str).tolist())
            if "component" in col.lower():
                seeds["components"].extend(df[col].dropna().astype(str).tolist())
    except Exception:
        pass
    return seeds

def chunk_text(text: str, max_tokens: int = 800, overlap: int = 100) -> List[str]:
    words = text.split()
    if not words:
        return []
    chunks = []
    i = 0
    while i < len(words):
        chunk = words[i:i + max_tokens]
        chunks.append(" ".join(chunk))
        i += max(1, max_tokens - overlap)
    return chunks

# --- Agent result dataclass ---
@dataclass
class AgentResult:
    name: str
    data: dict
    logs: List[str]

# --- Pattern-based extraction ---
_COMPONENT_PATTERNS = [
    r"\b(spindle|servo|stepper)\s+motor\b",
    r"\b(coolant|hydraulic)\s+pump\b",
    r"\b(ball|roller)\s+bearing\b",
    r"\b(tool\s+holder|collet|chuck)\b",
    r"\b(pressure|temperature|flow)\s+sensor\b",
    r"\b(limit\s+switch|proximity\s+switch)\b",
    r"\b(filter|nozzle|belt|pulley|seal|fuse|relay|valve|hose|manifold)\b",
]
_SUBSYSTEM_PATTERNS = [
    r"\b(hydraulic|lubrication|coolant|cooling|pneumatic)\s+system\b",
    r"\b(control\s+system|drive\s+system|safety\s+system)\b",
    r"\b(plc|pathpilot|hmi|scada)\b",
    r"\b(electrical\s+cabinet|spindle\s+drive|gearbox)\b",
]

def extract_candidates(chunks: List[str], top_k: int = 200) -> pd.DataFrame:
    """Candidate extraction from main.py"""
    vectorizer = TfidfVectorizer(ngram_range=(1, 3), min_df=1, max_df=0.95, stop_words="english")
    X = vectorizer.fit_transform(chunks)
    tfidf = np.asarray(X.mean(axis=0)).ravel()
    terms = np.array(vectorizer.get_feature_names_out())
    order = np.argsort(-tfidf)[:top_k]
    tfidf_terms = terms[order]

    # noun chunks via spaCy
    noun_terms = []
    nlp = get_nlp()
    if nlp is not None:
        for ch in chunks[:200]:
            try:
                doc = nlp(ch)
                for np_ in getattr(doc, "noun_chunks", []):
                    t = np_.text.strip().lower()
                    if 2 <= len(t) <= 60:
                        noun_terms.append(t)
            except Exception:
                continue

    # regex patterns
    regex_hits = []
    for ch in chunks:
        low = ch.lower()
        for pat in _COMPONENT_PATTERNS + _SUBSYSTEM_PATTERNS:
            for m in re2.finditer(pat, low, flags=re2.IGNORECASE):
                regex_hits.append(low[m.start():m.end()].strip())

    all_terms = list(tfidf_terms) + noun_terms + regex_hits
    cleaned = []
    for t in all_terms:
        t = re.sub(r"[^a-z0-9\-\s()/]", "", t.lower()).strip()
        t = re.sub(r"\s+", " ", t)
        if len(t) >= 3:
            cleaned.append(t)

    vc = pd.Series(cleaned).value_counts().reset_index()
    vc.columns = ["term", "raw_score"]
    vc["source_score"] = 0.0
    seed_subs = DEFAULT_SUBSYSTEM_SEEDS
    seed_comps = DEFAULT_COMPONENT_SEEDS
    for i, row in vc.iterrows():
        term = row["term"]
        boost = 0.0
        if any(re2.search(p, term) for p in _COMPONENT_PATTERNS):
            boost += 1.0
        if any(re2.search(p, term) for p in _SUBSYSTEM_PATTERNS):
            boost += 1.0
        if term in seed_subs or term in seed_comps:
            boost += 0.5
        vc.at[i, "source_score"] = boost
    vc["score"] = vc["raw_score"] + vc["source_score"]
    return vc

# --- Classification ---
@dataclass
class ClassifierConfig:
    use_llm: bool = False
    ollama_url: str = "http://localhost:11434/api/generate"
    ollama_model: str = "llama3.2:1b"
    timeout_s: int = 60

LABELS = ["Subsystem", "Component", "Irrelevant"]

_DEF_SYSTEM_PROMPT = (
    "You are a maintenance engineer. Classify each term as Subsystem, Component, or Irrelevant.\n"
    "Subsystems are higher-level systems (e.g., lubrication system, control cabinet).\n"
    "Components are parts (e.g., spindle motor, coolant pump, bearing).\n"
    "Return JSON: [{\"term\":..., \"label\":..., \"confidence\":0-1}] only."
)

def call_ollama(terms: List[str], cfg: ClassifierConfig) -> Optional[List[Dict]]:
    payload = {
        "model": cfg.ollama_model,
        "prompt": _DEF_SYSTEM_PROMPT + "\nTerms:\n" + "\n".join(f"- {t}" for t in terms) + "\n",
        "stream": False,
    }
    try:
        r = requests.post(cfg.ollama_url, json=payload, timeout=cfg.timeout_s)
        r.raise_for_status()
        text = r.json().get("response", "").strip()
        match = re.search(r"\[.*\]", text, re.S)
        if match:
            return json.loads(match.group(0))
        return json.loads(text)
    except Exception:
        return None

def heuristic_classify(terms: List[str]) -> List[Dict]:
    """Heuristic classification"""
    out = []
    for t in terms:
        low = t.lower()
        score_sub = score_cmp = 0.0
        if any(re2.search(p, low) for p in _SUBSYSTEM_PATTERNS):
            score_sub += 1.0
        if any(re2.search(p, low) for p in _COMPONENT_PATTERNS):
            score_cmp += 1.0
        if low in DEFAULT_SUBSYSTEM_SEEDS:
            score_sub += 0.5
        if low in DEFAULT_COMPONENT_SEEDS:
            score_cmp += 0.5
        if low.endswith(" system") or low in {"hmi", "plc", "scada"}:
            score_sub += 0.4
        if any(k in low for k in
               ["motor", "pump", "bearing", "sensor", "valve", "relay", "encoder", "filter", "nozzle", "belt", "seal"]):
            score_cmp += 0.4
        if score_sub == 0 and score_cmp == 0:
            label, conf = "Irrelevant", 0.35
        elif score_sub >= score_cmp:
            label, conf = "Subsystem", min(1.0, 0.5 + 0.3 * score_sub)
        else:
            label, conf = "Component", min(1.0, 0.5 + 0.3 * score_cmp)
        out.append({"term": t, "label": label, "confidence": round(conf, 3)})
    return out

def agent_classify(cands_df: pd.DataFrame, cfg: ClassifierConfig) -> AgentResult:
    """Classification agent"""
    logs = []
    df = cands_df.sort_values("score", ascending=False).head(300).copy()
    terms = df["term"].tolist()
    results = None
    if cfg.use_llm:
        logs.append("Calling Ollama classifier...")
        results = call_ollama(terms, cfg)
        if results is None:
            logs.append("LLM classification failed; falling back to heuristics.")
    if results is None:
        results = heuristic_classify(terms)
    res_df = pd.DataFrame(results)
    merged = df.merge(res_df, on="term", how="left")
    return AgentResult("classification", {"classified": merged}, logs)

# --- Normalization ---
def agent_normalize(df: pd.DataFrame, sim_threshold: float = 0.80) -> AgentResult:
    """Normalization agent"""
    logs = []
    em = get_sentence_embedder()
    labels, canonical = [], []
    if em is not None and len(df) > 0:
        emb = em.encode(df["term"].tolist(), normalize_embeddings=True)
        sim = np.matmul(emb, emb.T)
        assigned = [-1] * len(df)
        clusters = []
        for i in range(len(df)):
            if assigned[i] != -1:
                continue
            idxs = [j for j in range(len(df)) if sim[i, j] >= 0.86]
            for j in idxs: assigned[j] = len(clusters)
            clusters.append(idxs)
        for cl in clusters:
            reps = df.iloc[cl]["term"].tolist()
            can = min(reps, key=len)
            for j in cl:
                canonical.append(can)
                labels.append(j)
        logs.append(f"Embedding clusters: {len(set(canonical))} groups")
    else:
        terms = df["term"].tolist()
        used, can_map = set(), {}
        for i, t in enumerate(terms):
            if i in used: continue
            can = t
            used.add(i)
            for j in range(i + 1, len(terms)):
                if j in used: continue
                if fuzz.token_set_ratio(t, terms[j]) >= int(sim_threshold * 100):
                    used.add(j)
                    can_map[terms[j]] = can
            can_map[t] = can
        canonical = [can_map[t] for t in terms]
        logs.append(f"Fuzzy groups: {len(set(canonical))} groups")
    out = df.copy()
    out["canonical"] = canonical if len(canonical) == len(out) else out["term"].values
    agg = (out.groupby("canonical")
           .apply(lambda g: g.sort_values(["confidence", "score"], ascending=False).head(1))
           .reset_index(drop=True))
    return AgentResult("normalization", {"normalized": agg}, logs)

# --- Validation ---
def agent_validate(df: pd.DataFrame) -> AgentResult:
    """Validation agent"""
    logs = []
    df = df.copy()
    df["final_confidence"] = (
            0.5 * df.get("confidence", 0.5).fillna(0.5) +
            0.3 * (df["source_score"] / 2.0).clip(0, 1) +
            0.2 * (df["raw_score"].rank(pct=True))
    )
    keep = (df["label"] != "Irrelevant") & (df["final_confidence"] >= 0.45)
    filtered = df[keep].copy()
    logs.append(f"Validated: {len(filtered)}/{len(df)} kept")
    return AgentResult("validation", {"validated": filtered}, logs)

# --- Ingestion agent ---
def agent_ingest(files: List[Tuple[str, bytes]]) -> AgentResult:
    """Ingestion agent"""
    texts = []
    provenance = []
    logs = []
    for fname, blob in files:
        if fname.lower().endswith(".pdf"):
            txt = load_pdf(blob)
        elif fname.lower().endswith(".docx"):
            txt = load_docx(blob)
        elif fname.lower().endswith((".xlsx", "xls", "csv")):
            continue  # handled separately
        else:
            txt = load_txt(blob)
        logs.append(f"Loaded {fname}: {len(txt)} chars")
        provenance.append({"file": fname, "length": len(txt)})
        texts.append((fname, txt))
    full_text = "\n\n".join(t for _, t in texts)
    chunks = chunk_text(full_text)
    logs.append(f"Chunked into {len(chunks)} chunks")
    return AgentResult("ingestion", {"texts": texts, "chunks": chunks, "provenance": provenance, "full_text": full_text}, logs)

# --- Pydantic models for FMEA ---
class FMEAEntry(BaseModel):
    """Structured FMEA entry model"""
    component: str = Field(description="Component or system being analyzed")
    function: str = Field(description="Function of the component")
    failure_mode: str = Field(description="How the component can fail")
    potential_effects: str = Field(description="Effects of the failure")
    severity: int = Field(description="Severity rating (1-10)", ge=1, le=10)
    potential_causes: str = Field(description="Root causes of the failure")
    occurrence: int = Field(description="Occurrence rating (1-10)", ge=1, le=10)
    current_controls: str = Field(description="Current detection/prevention controls")
    detection: int = Field(description="Detection rating (1-10)", ge=1, le=10)
    rpn: int = Field(description="Risk Priority Number (Severity × Occurrence × Detection)")
    recommended_actions: str = Field(description="Recommended corrective actions")
    source_document: str = Field(description="Source document filename")
    page_section: str = Field(description="Page or section reference")

class FMEABatch(BaseModel):
    """Batch of FMEA entries"""
    entries: List[FMEAEntry] = Field(description="List of FMEA entries extracted")
    extraction_metadata: Dict[str, Any] = Field(description="Metadata about extraction")

# --- Enhanced Integrated Document Agent ---
class EnhancedIntegratedDocumentAgent:
    """
    Enhanced Integrated DocumentAgent combining multi-agent pipeline with LlamaIndex RAG
    Now includes intelligent FMEA generation from main2.py
    """
    
    def __init__(
        self,
        storage_dir: str = "./storage",
        use_local_llm: bool = True,
        local_model: str = "llama3.2:1b",
        embedding_model: str = "all-MiniLM-L6-v2", 
        chunk_size: int = 512,
        chunk_overlap: int = 50
    ):
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(exist_ok=True)
        
        # Initialize debug handler
        self.debug_handler = LlamaDebugHandler(print_trace_on_end=True)
        self.callback_manager = CallbackManager([self.debug_handler])
        
        # Initialize embedding model
        if _LANGCHAIN_OK:
            # Try local embedding path first, fallback to model name
            embedding_path = Path(__file__).parent.parent / "local_models" / "bge-small-en-v1.5"
            if embedding_path.exists():
                lc_hf = LC_HFEmbeddings(
                    model_name=str(embedding_path),
                    cache_folder=str(Path("./storage/embeddings")),
                    model_kwargs={"trust_remote_code": True}
                )
                self.embed_model = LangchainEmbedding(lc_hf)
            else:
                self.embed_model = HuggingFaceEmbedding(model_name=embedding_model)
        else:
            self.embed_model = HuggingFaceEmbedding(model_name=embedding_model)

        # Initialize LLM
        if use_local_llm:
            self.llm = Ollama(
                model=local_model,
                request_timeout=600.0,
                temperature=0.1,
                additional_kwargs={
                    "num_predict": 1024,
                    "top_p": 0.9,
                    "keep_alive": "30m"
                },
            )
        else:
            self.llm = OpenAI(
                model="gpt-4-turbo-preview",
                temperature=0.1
            )
        
        # Configure global settings
        Settings.llm = self.llm
        Settings.embed_model = self.embed_model
        Settings.callback_manager = self.callback_manager
        
        # Initialize node parser
        self.node_parser = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separator=" "
        )
        
        # Storage containers
        self.documents: List[Document] = []
        self.file_contents: List[Tuple[str, bytes]] = []  # For pipeline processing
        self.vector_index: Optional[VectorStoreIndex] = None
        self.bm25_retriever: Optional[BM25Retriever] = None
        self.query_engine: Optional[RetrieverQueryEngine] = None
        self.pipeline_results: Optional[Dict] = None  # Store pipeline results
        
        logger.info(f"EnhancedIntegratedDocumentAgent initialized with {'local' if use_local_llm else 'cloud'} LLM")

    def load_documents(self, document_paths: Union[str, List[str]]) -> None:
        """Load documents for both RAG and pipeline processing"""
        if isinstance(document_paths, str):
            document_paths = [document_paths]
        
        all_docs = []
        file_contents = []
        
        for doc_path in document_paths:
            path = Path(doc_path)
            
            if path.is_file():
                # Load file content for pipeline
                with open(path, 'rb') as f:
                    content = f.read()
                    file_contents.append((path.name, content))
                
                # Load with LlamaIndex for RAG
                reader = SimpleDirectoryReader(input_files=[str(path)])
            else:
                # Directory - load all files
                reader = SimpleDirectoryReader(input_dir=str(path))
                
                # Also load raw contents for pipeline
                for file_path in path.glob("*"):
                    if file_path.is_file() and file_path.suffix.lower() in {'.pdf', '.docx', '.txt', '.xlsx', '.xls', '.csv'}:
                        with open(file_path, 'rb') as f:
                            content = f.read()
                            file_contents.append((file_path.name, content))
            
            try:
                docs = reader.load_data()
                
                # Add metadata to documents
                for doc in docs:
                    doc.metadata.update({
                        'source_path': str(path),
                        'loaded_at': datetime.now().isoformat(),
                        'file_type': path.suffix.lower() if path.is_file() else 'directory'
                    })
                
                all_docs.extend(docs)
                logger.info(f"Loaded {len(docs)} documents from {path}")
                
            except Exception as e:
                logger.error(f"Failed to load documents from {path}: {e}")
        
        self.documents.extend(all_docs)
        self.file_contents.extend(file_contents)
        logger.info(f"Total documents loaded: {len(self.documents)}")

    def build_indexes(self) -> None:
        """Build indexes for RAG retrieval"""
        if not self.documents:
            raise ValueError("No documents loaded. Call load_documents() first.")
        
        logger.info("Building indexes...")
        
        # Parse documents into nodes
        nodes = self.node_parser.get_nodes_from_documents(self.documents)
        logger.info(f"Created {len(nodes)} nodes from documents")
        
        # Build FAISS vector store
        dimension = 384  # Default embedding dimension
        faiss_index = faiss.IndexFlatIP(dimension)
        vector_store = FaissVectorStore(faiss_index=faiss_index)
        storage_context = StorageContext.from_defaults(vector_store=vector_store)
        
        # Build vector index
        self.vector_index = VectorStoreIndex(
            nodes=nodes,
            storage_context=storage_context
        )
        
        # Build BM25 retriever
        self.bm25_retriever = BM25Retriever.from_defaults(
            nodes=nodes,
            similarity_top_k=5
        )
        
        # Create hybrid retriever
        vector_retriever = self.vector_index.as_retriever(similarity_top_k=5)
        
        self.fusion_retriever = QueryFusionRetriever(
            retrievers=[self.bm25_retriever, vector_retriever],
            similarity_top_k=10,
            num_queries=2
        )
        
        # Build query engine
        self.query_engine = RetrieverQueryEngine.from_args(
            retriever=self.fusion_retriever,
            response_mode="compact",
            verbose=True
        )
        
        logger.info("Indexes built successfully")

    def run_extraction_pipeline(
        self,
        top_k: int = 300,
        sim_threshold: float = 0.80,
        use_llm: bool = False,
        ollama_url: str = "http://localhost:11434/api/generate",
        ollama_model: str = "llama3.2:1b"
    ) -> Dict:
        """Run the multi-agent extraction pipeline"""
        if not self.file_contents:
            raise ValueError("No file contents loaded. Call load_documents() first.")
        
        logger.info("Running extraction pipeline...")
        pipeline_logs = []
        
        # Extract seeds from Excel/CSV files
        custom_seeds = {"subsystems": [], "components": []}
        for fname, blob in self.file_contents:
            if fname.lower().endswith((".xlsx", ".xls", ".csv")):
                seeds = load_excel_csv(blob, fname)
                custom_seeds["subsystems"].extend(seeds["subsystems"])
                custom_seeds["components"].extend(seeds["components"])
        
        # Merge with default seeds
        DEFAULT_SUBSYSTEM_SEEDS.update(s.lower() for s in custom_seeds["subsystems"])
        DEFAULT_COMPONENT_SEEDS.update(c.lower() for c in custom_seeds["components"])
        
        # 1. Ingestion
        start = time.time()
        ing = agent_ingest(self.file_contents)
        pipeline_logs.extend([f"[Ingest] {l}" for l in ing.logs])
        logger.info(f"Ingestion complete: {len(ing.data['chunks'])} chunks ({time.time() - start:.1f}s)")
        
        # 2. Candidate extraction
        start = time.time()
        cands = extract_candidates(ing.data['chunks'], top_k=top_k)
        logger.info(f"Candidate extraction complete: {len(cands)} terms ({time.time() - start:.1f}s)")
        
        # 3. Classification
        start = time.time()
        cfg = ClassifierConfig(
            use_llm=use_llm,
            ollama_url=ollama_url,
            ollama_model=ollama_model
        )
        cls = agent_classify(cands, cfg)
        pipeline_logs.extend([f"[Classify] {l}" for l in cls.logs])
        logger.info(f"Classification complete ({time.time() - start:.1f}s)")
        
        # 4. Normalization
        start = time.time()
        norm = agent_normalize(cls.data['classified'], sim_threshold=sim_threshold)
        pipeline_logs.extend([f"[Normalize] {l}" for l in norm.logs])
        logger.info(f"Normalization complete ({time.time() - start:.1f}s)")
        
        # 5. Validation
        start = time.time()
        val = agent_validate(norm.data['normalized'])
        pipeline_logs.extend([f"[Validate] {l}" for l in val.logs])
        logger.info(f"Validation complete: {len(val.data['validated'])} kept ({time.time() - start:.1f}s)")
        
        # Store results
        self.pipeline_results = {
            "ingestion": ing.data,
            "candidates": cands,
            "classified": cls.data['classified'],
            "normalized": norm.data['normalized'],
            "validated": val.data['validated'],
            "logs": pipeline_logs,
            "extraction_metadata": {
                "extracted_at": datetime.now().isoformat(),
                "total_files": len(self.file_contents),
                "total_chunks": len(ing.data['chunks']),
                "total_candidates": len(cands),
                "pipeline_config": {
                    "top_k": top_k,
                    "sim_threshold": sim_threshold,
                    "use_llm": use_llm,
                    "ollama_model": ollama_model if use_llm else None
                }
            }
        }
        
        logger.info("Extraction pipeline completed successfully")
        return self.pipeline_results

    def generate_intelligent_fmea_table(self, components_df: pd.DataFrame) -> pd.DataFrame:
        """Generate enhanced FMEA table with intelligent analysis from main2.py"""
        if not self.pipeline_results or 'ingestion' not in self.pipeline_results:
            # Fallback to basic historical text if no pipeline results
            historical_text = ""
        else:
            historical_text = self.pipeline_results['ingestion'].get('full_text', '')
        
        # Use the enhanced FMEA generation from main2.py
        fmea_table = generate_intelligent_fmea(components_df, historical_text)
        
        logger.info(f"Generated intelligent FMEA table with {len(fmea_table)} entries")
        return fmea_table

    def chat_with_documents(self, question: str) -> Dict[str, Any]:
        """RAG-based chat interface with source citation"""
        if not self.query_engine:
            raise ValueError("Indexes not built. Call build_indexes() first.")
        
        logger.info(f"Processing chat question: {question}")
        
        # Query the engine
        response = self.query_engine.query(question)
        
        # Extract source information
        sources = []
        for node in response.source_nodes:
            source_info = {
                'content': node.node.text[:200] + "..." if len(node.node.text) > 200 else node.node.text,
                'score': node.score,
                'metadata': node.node.metadata,
                'source_document': node.node.metadata.get('file_name', 'Unknown'),
                'page_info': node.node.metadata.get('page_label', 'N/A')
            }
            sources.append(source_info)
        
        return {
            'answer': str(response),
            'sources': sources,
            'question': question,
            'timestamp': datetime.now().isoformat(),
            'num_sources': len(sources)
        }

    def get_document_summary(self) -> Dict[str, Any]:
        """Get summary of loaded documents"""
        if not self.documents:
            return {"message": "No documents loaded"}
        
        summary = {
            'total_documents': len(self.documents),
            'documents': [],
            'total_characters': 0,
            'file_types': {}
        }
        
        for doc in self.documents:
            doc_info = {
                'filename': doc.metadata.get('file_name', 'Unknown'),
                'file_type': doc.metadata.get('file_type', 'Unknown'),
                'character_count': len(doc.text),
                'loaded_at': doc.metadata.get('loaded_at', 'Unknown')
            }
            summary['documents'].append(doc_info)
            summary['total_characters'] += doc_info['character_count']
            
            # Count file types
            file_type = doc_info['file_type']
            summary['file_types'][file_type] = summary['file_types'].get(file_type, 0) + 1
        
        return summary

    def get_pipeline_results(self) -> Optional[Dict]:
        """Get the results from the extraction pipeline"""
        return self.pipeline_results

    def save_results(self, output_path: str) -> None:
        """Save pipeline results to JSON file"""
        if not self.pipeline_results:
            raise ValueError("No pipeline results available. Run extraction pipeline first.")
        
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Convert DataFrames to dictionaries for JSON serialization
        serializable_results = {}
        for key, value in self.pipeline_results.items():
            if isinstance(value, pd.DataFrame):
                serializable_results[key] = value.to_dict('records')
            else:
                serializable_results[key] = value
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(serializable_results, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Pipeline results saved to {output_file}")


# Example usage and testing
def main():
    """Example usage of the EnhancedIntegratedDocumentAgent"""
    
    # Initialize the agent
    agent = EnhancedIntegratedDocumentAgent(
        storage_dir="./enhanced_integrated_storage",
        use_local_llm=True,
        local_model="llama3.2:1b"
    )
    
    # Load documents 
    document_paths = [
        "./sample_documents/",  # Directory with documents
        
    ]
    
    try:
        # Load documents
        agent.load_documents(document_paths)
        
        # Get document summary
        summary = agent.get_document_summary()
        print("Document Summary:", summary)
        
        # Run extraction pipeline
        pipeline_results = agent.run_extraction_pipeline(
            top_k=300,
            sim_threshold=0.80,
            use_llm=False  # Use heuristic classification
        )
        
        print(f"\nPipeline Results:")
        print(f"- Total validated terms: {len(pipeline_results['validated'])}")
        
        # Show some results
        validated_df = pipeline_results['validated']
        components_df = validated_df[validated_df['label'] == 'Component']
        subsystems_df = validated_df[validated_df['label'] == 'Subsystem']
        
        print(f"- Components found: {len(components_df)}")
        print(f"- Subsystems found: {len(subsystems_df)}")
        
        # Generate enhanced FMEA table
        if not components_df.empty:
            fmea_table = agent.generate_intelligent_fmea_table(components_df)
            print(f"\nGenerated enhanced FMEA table with {len(fmea_table)} entries")
            
            # Show sample of enhanced FMEA data
            print("\nSample FMEA entries:")
            for idx, row in fmea_table.head(3).iterrows():
                print(f"\nComponent: {row['Component']}")
                print(f"Failure Mode: {row['Potential Failure Mode']}")
                print(f"Severity: {row['Severity (S)']} ({row['Severity Level']}) - {row['Severity Reasoning']}")
                print(f"RPN: {row['RPN']}")
                print(f"Data Source: {row['Data Source']}")
                print(f"Recommendations: {row['Detailed Recommended Actions'][:100]}...")
        
        # Build indexes for RAG
        agent.build_indexes()
        
        # Demo chat interface
        chat_questions = [
            "What components were found in the documents?",
            "What are the main subsystems identified?",
            "Which components might have high failure risk?",
            "What maintenance recommendations can you suggest?"
        ]
        
        print("\n" + "="*50)
        print("CHAT WITH DOCUMENTS DEMO")
        print("="*50)
        
        for question in chat_questions:
            print(f"\nQ: {question}")
            try:
                response = agent.chat_with_documents(question)
                print(f"A: {response['answer']}")
                print(f"Sources: {response['num_sources']} document chunks")
            except Exception as e:
                print(f"Chat error: {e}")
        
        # Save results
        agent.save_results("./results/enhanced_integrated_extraction.json")
        print("\nResults saved to ./results/enhanced_integrated_extraction.json")
    
    except Exception as e:
        logger.error(f"Error in main execution: {e}")
        raise


if __name__ == "__main__":
    main()
